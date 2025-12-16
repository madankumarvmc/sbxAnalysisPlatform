"""
MS Word Report Generator for Warehouse Analysis Tool

This module generates comprehensive Word reports with:
- LLM-powered insights
- Tables from analysis results  
- Charts and visualizations
- Professional formatting
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import pandas as pd
import io
from datetime import datetime
from typing import Dict, Any, Optional

from .gemini_client import GeminiClient
from .word_prompts import PROMPTS, FALLBACKS
from .chart_recreator import (
    create_receipt_volume_chart,
    create_order_volume_chart, 
    create_abc_distribution_chart,
    create_percentile_chart
)


class WordReportGenerator:
    """
    Generate MS Word reports from warehouse analysis results.
    """
    
    def __init__(self, analysis_results: Dict[str, Any], config: Optional[Dict[str, Any]] = None):
        """
        Initialize the Word report generator.
        
        Args:
            analysis_results: Complete analysis results from all modules
            config: Configuration settings (analysis variables)
        """
        self.analysis_results = analysis_results
        self.config = config or {}
        self.doc = Document()
        
        # Initialize Gemini client (will handle errors gracefully)
        try:
            self.gemini = GeminiClient()
            self.llm_available = self.gemini.test_connection()
        except Exception as e:
            self.gemini = None
            self.llm_available = False
        
        # Set document properties
        self._setup_document_styles()
    
    def _setup_document_styles(self):
        """Setup document-wide styles and formatting."""
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
    
    def generate_report(self) -> io.BytesIO:
        """
        Generate the complete Word report.
        
        Returns:
            BytesIO buffer containing the Word document
        """
        try:
            # Add sections in predefined order
            self.add_title_page()
            self.add_executive_summary()
            
            # Add Receipt Analysis if available
            if 'receipt_analysis' in self.analysis_results:
                receipt_data = self.analysis_results['receipt_analysis']
                # Check if it's a successful analysis result
                if isinstance(receipt_data, dict):
                    if receipt_data.get('success', True):
                        self.add_receipt_analysis()
                else:
                    # If it's not a dict, assume it's valid analysis data
                    self.add_receipt_analysis()
            
            # Future sections (to be added)
            # self.add_order_analysis()
            # self.add_sku_analysis()
            # self.add_inventory_analysis()
            # self.add_manpower_analysis()
            
            # Save to buffer
            doc_buffer = io.BytesIO()
            self.doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer
            
        except Exception as e:
            raise
    
    def add_title_page(self):
        """Add professional title page."""
        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('WAREHOUSE ANALYSIS REPORT')
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Subtitle
        self.doc.add_paragraph()
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Comprehensive Operational Analysis')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(68, 114, 196)
        
        # Add spacing
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Report metadata
        metadata = self.doc.add_paragraph()
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        run = metadata.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}\n')
        run.font.size = Pt(12)
        
        # Analysis period (if available)
        if 'data_loader' in self.analysis_results:
            data_info = self.analysis_results['data_loader']
            if 'data' in data_info:
                run = metadata.add_run(f'Data Source: Warehouse Analytics System\n')
                run.font.size = Pt(12)
        
        # Add spacing before content
        for _ in range(5):
            self.doc.add_paragraph()
        
        # Disclaimer
        disclaimer = self.doc.add_paragraph()
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = disclaimer.add_run('This report contains confidential operational data and analysis')
        run.font.size = Pt(10)
        run.font.italic = True
        
        # Page break
        self.doc.add_page_break()
    
    def add_executive_summary(self):
        """Add executive summary section with LLM insights."""
        self.doc.add_heading('Executive Summary', level=1)
        
        # Gather all key metrics for the summary
        all_metrics = self._gather_key_metrics()
        
        # Generate LLM insights
        if self.llm_available and self.gemini:
            prompt = PROMPTS['executive_summary']['main'].format(
                all_metrics=self.gemini.format_data_for_prompt(all_metrics)
            )
            insights = self.gemini.generate_insight(
                prompt,
                FALLBACKS['executive_summary']
            )
        else:
            insights = FALLBACKS['executive_summary']
        
        # Add insights as paragraphs
        for paragraph in insights.split('\n\n'):
            if paragraph.strip():
                p = self.doc.add_paragraph(paragraph.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Add key metrics table
        self.doc.add_heading('Key Performance Indicators', level=2)
        self._add_kpi_table(all_metrics)
        
        # Page break
        self.doc.add_page_break()
    
    def add_receipt_analysis(self):
        """Add receipt analysis section with insights, tables, and charts."""
        self.doc.add_heading('Receipt Analysis', level=1)
        
        receipt_data = self.analysis_results.get('receipt_analysis', {})
        
        # Section 1: Daily Receipt Pattern Analysis
        daily_data_key = 'daily_patterns' if 'daily_patterns' in receipt_data else 'daily_summary'
        if daily_data_key in receipt_data:
            self.doc.add_heading('Daily Receipt Patterns', level=2)
            
            daily_data = receipt_data[daily_data_key]
            
            # Generate LLM insights (before table)
            if self.llm_available and self.gemini:
                prompt = PROMPTS['receipt_analysis']['daily_summary'].format(
                    daily_data=self.gemini.format_data_for_prompt(daily_data)
                )
                insights = self.gemini.generate_insight(
                    prompt,
                    FALLBACKS['receipt_analysis']
                )
            else:
                insights = FALLBACKS['receipt_analysis']
            
            # Add insights
            p = self.doc.add_paragraph(insights)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add daily patterns table
            self.add_dataframe_as_table(
                daily_data,
                'Daily Receipt Patterns'
            )
            
            # Add chart
            try:
                chart_buffer = create_receipt_volume_chart(daily_data)
                self.doc.add_picture(chart_buffer, width=Inches(6))
                last_paragraph = self.doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                pass  # Skip chart if creation fails
        
        # Section 2: Percentile Analysis
        if 'percentile_analysis' in receipt_data:
            self.doc.add_heading('Receipt Volume Percentile Analysis', level=2)
            
            # Generate insights for percentiles
            if self.llm_available and self.gemini:
                prompt = PROMPTS['receipt_analysis']['percentile_analysis'].format(
                    percentile_data=self.gemini.format_data_for_prompt(receipt_data['percentile_analysis'])
                )
                percentile_insights = self.gemini.generate_insight(
                    prompt,
                    "Percentile analysis shows volume distribution for capacity planning."
                )
                
                p = self.doc.add_paragraph(percentile_insights)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add percentile table
            self.add_dataframe_as_table(
                receipt_data['percentile_analysis'],
                'Volume at Different Percentiles'
            )
            
            # Add percentile chart if possible
            try:
                chart_buffer = create_percentile_chart(
                    receipt_data['percentile_analysis'],
                    'Receipt Volume Percentiles'
                )
                self.doc.add_picture(chart_buffer, width=Inches(6))
                last_paragraph = self.doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                pass  # Skip chart if creation fails
        
        # Section 3: Overall Recommendations
        summary_key = 'data_summary' if 'data_summary' in receipt_data else 'summary_statistics'
        if summary_key in receipt_data:
            self.doc.add_heading('Receipt Operations Recommendations', level=2)
            
            if self.llm_available and self.gemini:
                prompt = PROMPTS['receipt_analysis']['overall'].format(
                    summary_data=self.gemini.format_data_for_prompt(receipt_data[summary_key])
                )
                recommendations = self.gemini.generate_insight(
                    prompt,
                    "Optimize receiving operations based on identified patterns."
                )
            else:
                recommendations = "Optimize receiving operations based on identified patterns."
                
            p = self.doc.add_paragraph(recommendations)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Page break
        self.doc.add_page_break()
    
    def add_dataframe_as_table(self, df: pd.DataFrame, title: Optional[str] = None):
        """
        Convert pandas DataFrame to Word table with default styling.
        
        Args:
            df: DataFrame to convert
            title: Optional title for the table
        """
        if title:
            heading = self.doc.add_heading(title, level=3)
        
        # Handle empty DataFrame
        if df.empty:
            self.doc.add_paragraph("No data available for this analysis.")
            return
        
        # Create table with headers
        table = self.doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light List Accent 1'  # Word default table style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            header_cells[i].text = str(col)
            # Bold headers
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add data rows (limit to reasonable number for Word doc)
        max_rows = min(len(df), 100)  # Limit to 100 rows for readability
        
        for idx, row in df.head(max_rows).iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                # Format values appropriately
                if pd.isna(value):
                    cell_text = ''
                elif isinstance(value, (int, float)):
                    if abs(value) >= 1000:
                        cell_text = f'{value:,.0f}'
                    elif isinstance(value, float):
                        cell_text = f'{value:.2f}'
                    else:
                        cell_text = str(value)
                else:
                    cell_text = str(value)
                
                row_cells[i].text = cell_text
        
        # Add note if data was truncated
        if len(df) > max_rows:
            note = self.doc.add_paragraph()
            note.add_run(f'Note: Showing {max_rows} of {len(df)} total rows').italic = True
        
        # Add spacing after table
        self.doc.add_paragraph()
    
    def _gather_key_metrics(self) -> Dict[str, Any]:
        """
        Gather all key metrics from analysis results for executive summary.
        
        Returns:
            Dictionary of key metrics
        """
        metrics = {
            'analysis_date': datetime.now().strftime('%Y-%m-%d'),
            'sections_completed': []
        }
        
        # Order Analysis Metrics
        if 'order_analysis' in self.analysis_results:
            order_data = self.analysis_results['order_analysis']
            if isinstance(order_data, dict) and order_data.get('success', True):
                metrics['sections_completed'].append('Order Analysis')
                # Look for multiple possible summary keys
                for key in ['summary_statistics', 'data_summary', 'statistics']:
                    if key in order_data:
                        metrics['order_metrics'] = order_data[key]
                        break
        
        # Receipt Analysis Metrics  
        if 'receipt_analysis' in self.analysis_results:
            receipt_data = self.analysis_results['receipt_analysis']
            if isinstance(receipt_data, dict) and receipt_data.get('success', True):
                metrics['sections_completed'].append('Receipt Analysis')
                # Look for multiple possible summary keys
                for key in ['data_summary', 'summary_statistics', 'statistics']:
                    if key in receipt_data:
                        metrics['receipt_metrics'] = receipt_data[key]
                        break
        
        # SKU Analysis Metrics
        if 'sku_analysis' in self.analysis_results:
            sku_data = self.analysis_results['sku_analysis']
            if isinstance(sku_data, dict) and sku_data.get('success', True):
                metrics['sections_completed'].append('SKU Analysis')
                for key in ['summary', 'data_summary', 'statistics']:
                    if key in sku_data:
                        metrics['sku_metrics'] = sku_data[key]
                        break
        
        # ABC-FMS Analysis Metrics
        if 'abc_fms_analysis' in self.analysis_results:
            abc_data = self.analysis_results['abc_fms_analysis']
            if isinstance(abc_data, dict) and abc_data.get('success', True):
                metrics['sections_completed'].append('ABC-FMS Analysis')
                for key in ['summary', 'data_summary', 'statistics']:
                    if key in abc_data:
                        metrics['abc_fms_metrics'] = abc_data[key]
                        break
        
        # Inventory Analysis Metrics
        if 'inventory_analysis' in self.analysis_results:
            inv_data = self.analysis_results['inventory_analysis']
            if isinstance(inv_data, dict) and inv_data.get('success', True):
                metrics['sections_completed'].append('Inventory Analysis')
                for key in ['summary', 'data_summary', 'statistics']:
                    if key in inv_data:
                        metrics['inventory_metrics'] = inv_data[key]
                        break
        
        # Manpower Analysis Metrics
        if 'manpower_analysis' in self.analysis_results:
            manpower_data = self.analysis_results['manpower_analysis']
            if isinstance(manpower_data, dict) and manpower_data.get('success', True):
                metrics['sections_completed'].append('Manpower Analysis')
                for key in ['summary', 'data_summary', 'statistics']:
                    if key in manpower_data:
                        metrics['manpower_metrics'] = manpower_data[key]
                        break
        
        metrics['total_sections'] = len(metrics['sections_completed'])
        
        return metrics
    
    def _add_kpi_table(self, metrics: Dict[str, Any]):
        """
        Add a KPI summary table to the document.
        
        Args:
            metrics: Dictionary of key metrics
        """
        # Create KPI data
        kpi_data = []
        
        # Add available KPIs
        if 'order_metrics' in metrics:
            order_m = metrics['order_metrics']
            if isinstance(order_m, dict):
                kpi_data.append(['Daily Order Average', f"{order_m.get('avg_daily_orders', 'N/A'):,.0f}"])
                kpi_data.append(['Peak Order Volume', f"{order_m.get('max_daily_orders', 'N/A'):,.0f}"])
        
        if 'receipt_metrics' in metrics:
            receipt_m = metrics['receipt_metrics']
            if isinstance(receipt_m, dict):
                kpi_data.append(['Daily Receipt Average', f"{receipt_m.get('avg_daily_receipts', 'N/A'):,.0f}"])
                kpi_data.append(['Receipt Efficiency', f"{receipt_m.get('efficiency', 'N/A')}%"])
        
        if 'inventory_metrics' in metrics:
            inv_m = metrics['inventory_metrics']
            if isinstance(inv_m, dict):
                kpi_data.append(['Average Inventory Turnover', f"{inv_m.get('avg_turnover', 'N/A'):.1f}"])
                kpi_data.append(['Total Inventory Value', f"${inv_m.get('total_value', 0):,.0f}"])
        
        if 'manpower_metrics' in metrics:
            man_m = metrics['manpower_metrics']
            if isinstance(man_m, dict):
                kpi_data.append(['Required Staff (Peak)', f"{man_m.get('peak_staff_required', 'N/A')}"])
                kpi_data.append(['Current Efficiency', f"{man_m.get('current_efficiency', 'N/A')}%"])
        
        # Add sections completed
        kpi_data.append(['Analysis Sections Completed', f"{metrics['total_sections']}"])
        
        if kpi_data:
            # Create table
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            # Headers
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Key Performance Indicator'
            header_cells[1].text = 'Value'
            
            # Make headers bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add data rows
            for kpi_name, kpi_value in kpi_data:
                row_cells = table.add_row().cells
                row_cells[0].text = str(kpi_name)
                row_cells[1].text = str(kpi_value)
            
            # Add spacing
            self.doc.add_paragraph()